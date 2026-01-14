#Versión idem app9 con boton descarga WORD (Funciona bien corregido el 9G)
#carga doc + feedback + logoPNG + comentario feedback + Descarga Word

from __future__ import annotations

import os
import io
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from PIL import Image

# Local .env
load_dotenv()

# --- CONFIGURACIÓN DE AVATARES ---
try:
    logo_path = "assistant_logo.png"
    logo_path_user = "user_logo.png"
    USER_AVATAR = Image.open(logo_path_user)
    ASSISTANT_AVATAR = Image.open(logo_path)
except Exception:
    USER_AVATAR = "👤"
    ASSISTANT_AVATAR = "⚖️"


API_KEY = os.getenv("OPENAI_API_KEY", "")
VS_ID = os.getenv("OPENAI_VECTOR_STORE_ID", "")
MODEL = os.getenv("OPENAI_MODEL", "gpt-5.1")

st.set_page_config(page_title="KA Legal (Vector Store OpenAI)", layout="wide", initial_sidebar_state="expanded")

with st.sidebar:
    st.title("📂 Documentos")
    archivo = st.file_uploader("Cargar PDF o Word para analizar", type=["pdf", "docx"])
    texto_del_archivo = ""
    
    if archivo:
        if archivo.type == "application/pdf":
            reader = PdfReader(archivo)
            texto_del_archivo = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        else:
            doc = Document(archivo)
            texto_del_archivo = "\n".join([p.text for p in doc.paragraphs])
        st.success("Documento cargado con éxito")

col_logo, col_titulo = st.columns([1, 8]) # El 1 es para el logo (estrecho) y el 8 para el título (ancho)
with col_logo:
    st.image("logo.png", width=80) # Un tamaño más pequeño suele quedar mejor al lado del texto

with col_titulo:
    st.title("Chalk Legal")

if not API_KEY or not VS_ID:
    st.error("Faltan credenciales.")
    st.stop()

client = OpenAI(api_key=API_KEY)

if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system",
            "content": (
                "Sos un asistente legal societario (Argentina). "
                "Respondé en español, con citas cuando existan en los documentos. "
                "Si no hay soporte documental suficiente, decilo claramente."
                "Responder solo en base a los documentos que tienes cargados en VS."
            ),
        }
    ]

# Render de historial y feedback
for i, m in enumerate(st.session_state.messages):
    if m["role"] in ("user", "assistant"):
        avatar_to_use = USER_AVATAR if m["role"] == "user" else ASSISTANT_AVATAR
        with st.chat_message(m["role"], avatar=avatar_to_use):
            # Usamos display_content para la UI si existe
            content_to_show = m.get("display_content", m["content"])
            st.markdown(content_to_show)
            
            if m["role"] == "assistant" and i > 0:
                col1, col2 = st.columns([1, 1])
                with col1:
                    key_fb = f"feedback_{i}"
                    fback = st.feedback("thumbs", key=key_fb)
                with col2:
                    doc_download = Document()
                    doc_download.add_heading('Chalk Legal powered by TCA', 0)
                    doc_download.add_paragraph(m["content"])
                    buffer = io.BytesIO()
                    doc_download.save(buffer)
                    word_data = buffer.getvalue()
                    st.download_button(
                        label="📥 ⬇ Descargar a Word",
                        data=word_data,
                        file_name=f"chalk_legal_download_{i}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_{i}"
                    )
                if fback == 0:
                    with st.container():
                        comentario = st.text_area("¿Cómo podemos mejorar?", key=f"msg_{key_fb}")
                        if st.button("Enviar comentario", key=f"btn_{key_fb}"):
                            st.success("¡Gracias!")

user_text = st.chat_input("Escribí tu consulta legal…")

if user_text:
    prompt_final = user_text
    if texto_del_archivo:
        prompt_final = f"CONTENIDO DEL DOCUMENTO A ANALIZAR:\n{texto_del_archivo}\n\nCONSULTA: {user_text}"

    # Guardamos ambos para la visualización limpia
    st.session_state.messages.append({
        "role": "user", 
        "content": prompt_final, 
        "display_content": user_text
    })
    st.rerun()

# Lógica de respuesta
if len(st.session_state.messages) > 0 and st.session_state.messages[-1]["role"] == "user":
    with st.chat_message("assistant", avatar=ASSISTANT_AVATAR):
        placeholder = st.empty()
        
        # --- FILTRADO DE MENSAJES (Para evitar el error display_content) ---
        api_messages = [
            {"role": m["role"], "content": m["content"]} 
            for m in st.session_state.messages
        ]

        with st.spinner("Analizando documentos y leyes..."):
            # Volvemos a la estructura original que te funcionaba con Vector Stores
            resp = client.responses.create(
                model=MODEL,
                input=api_messages,
                tools=[{"type": "file_search", "vector_store_ids": [VS_ID]}],
            )

        answer_text = resp.output_text or "(sin respuesta)"
        placeholder.markdown(answer_text)
        st.session_state.messages.append({"role": "assistant", "content": answer_text})
        st.rerun()